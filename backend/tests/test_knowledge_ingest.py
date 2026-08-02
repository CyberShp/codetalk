from __future__ import annotations

from io import BytesIO

from docx import Document
from openpyxl import Workbook

from app.services.knowledge_ingest import MAX_SOURCE_BYTES, parse_bytes, parse_paste


def _text_pdf(text: str) -> bytes:
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("latin-1")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
    ]
    result = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, body in enumerate(objects, start=1):
        offsets.append(len(result))
        result.extend(f"{index} 0 obj\n".encode("ascii"))
        result.extend(body)
        result.extend(b"\nendobj\n")
    startxref = len(result)
    result.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    result.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        result.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    result.extend(f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{startxref}\n%%EOF\n".encode("ascii"))
    return bytes(result)


def test_paste_parser_keeps_immutable_snapshot_and_line_locators():
    parsed = parse_paste("first line\nsecond line\n")

    assert parsed.status == "parsed"
    assert parsed.text == "first line\nsecond line\n"
    assert [locator["kind"] for locator in parsed.locators] == ["line", "line"]
    assert parsed.locators[1]["start"] == 2


def test_docx_parser_emits_paragraph_and_table_cell_locators():
    document = Document()
    document.add_heading("Protocol", level=1)
    document.add_paragraph("Login state transition")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "condition"
    table.cell(0, 1).text = "login_complete=false"
    buffer = BytesIO()
    document.save(buffer)

    parsed = parse_bytes(buffer.getvalue(), filename="incident.docx")

    assert parsed.status == "parsed"
    assert any(locator["kind"] == "paragraph" and locator["heading_path"] == ["Protocol"] for locator in parsed.locators)
    assert any(locator["kind"] == "table-cell" and locator["cell"] == "B1" for locator in parsed.locators)


def test_text_pdf_parser_emits_page_text_run_and_scanned_pdf_needs_ocr():
    parsed = parse_bytes(_text_pdf("CmdSN window shrinks"), filename="incident.pdf")
    scanned = parse_bytes(_text_pdf(""), filename="scan.pdf")

    assert parsed.status == "parsed"
    assert parsed.locators[0]["kind"] == "page-text-run"
    assert parsed.locators[0]["page"] == 1
    assert scanned.status == "needs_ocr"


def test_xlsx_parser_emits_sheet_cell_locators():
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Incidents"
    sheet["A1"] = "module"
    sheet["B1"] = "symptom"
    sheet["A2"] = "dtoe"
    sheet["B2"] = "window update slow"
    buffer = BytesIO()
    workbook.save(buffer)

    parsed = parse_bytes(buffer.getvalue(), filename="incidents.xlsx")

    assert parsed.status == "parsed"
    assert any(locator["kind"] == "sheet-cell" and locator["sheet"] == "Incidents" and locator["cell"] == "B2" and locator["header"] == "symptom" for locator in parsed.locators)


def test_malformed_binary_source_returns_retryable_failed_parse_result():
    parsed = parse_bytes(b"not a pdf", filename="broken.pdf")

    assert parsed.status == "failed"
    assert parsed.parser == "pypdf"
    assert parsed.error


def test_ingest_limits_input_and_strips_links_unless_explicitly_retained():
    too_large = parse_bytes(b"x" * (MAX_SOURCE_BYTES + 1), filename="too-large.txt")
    stripped = parse_paste("see https://codehub.example/storage/array", keep_links=False)
    retained = parse_paste("see https://codehub.example/storage/array", keep_links=True)

    assert too_large.status == "failed"
    assert too_large.error == "source_too_large"
    assert "https://" not in stripped.text
    assert "https://" in retained.text
