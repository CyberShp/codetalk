"""Export service -- package reports into zip/docx/xml formats."""

import io
import logging
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import NamedTuple

from app.config import settings
from app.services.external_agent_discovery import redact_agent_diagnostic_text

logger = logging.getLogger(__name__)


class _ReportDoc(NamedTuple):
    name: str     # filename (e.g. "summary.md") — used in zip entry / xml filename attr
    content: str  # raw text content


# ---------------------------------------------------------------------------
# Public: task export
# ---------------------------------------------------------------------------

async def export_reports(task_id: str, fmt: str) -> tuple[bytes, str, str]:
    """Export task reports in the requested format.

    Returns:
        Tuple of (file_bytes, filename, content_type).

    Raises:
        FileNotFoundError: If no output files exist for the task.
        ValueError: If format is not supported.
    """
    output_dir = settings.outputs_path / task_id
    if not output_dir.exists():
        raise FileNotFoundError(f"任务输出目录不存在: {task_id}")

    md_files = sorted(output_dir.glob("*.md"))
    if not md_files:
        raise FileNotFoundError(f"任务无输出文件: {task_id}")

    docs = [_ReportDoc(name=f.name, content=f.read_text(encoding="utf-8")) for f in md_files]
    return _dispatch(docs, f"codetalk-{task_id[:8]}", fmt)


# ---------------------------------------------------------------------------
# Public: workspace report export
# ---------------------------------------------------------------------------

async def export_workspace_reports(
    ws_id: str, fmt: str, db: object, task_id: str | None = None
) -> tuple[bytes, str, str]:
    """Export completed workspace reports in the requested format.

    Args:
        ws_id: Workspace UUID.
        fmt: Export format — "md", "docx", or "xml".
        db: aiosqlite.Connection.
        task_id: Optional analysis task id; defaults to the latest completed
            task so workspace exports do not bundle historical reports.

    Raises:
        FileNotFoundError: No completed reports for this workspace.
        ValueError: Unsupported format.
    """
    resolved_task_id = task_id
    legacy_only = resolved_task_id == "__legacy__"

    if resolved_task_id is None:
        async with db.execute(  # type: ignore[attr-defined]
            "SELECT task_id FROM workspace_reports"
            " WHERE workspace_id = ? AND status = 'completed'"
            " ORDER BY created_at DESC LIMIT 1",
            (ws_id,),
        ) as cur:
            latest = await cur.fetchone()
        if latest:
            resolved_task_id = latest["task_id"]
            legacy_only = resolved_task_id is None

    where = "workspace_id = ? AND status = 'completed'"
    params: list[object] = [ws_id]
    if legacy_only:
        where += " AND task_id IS NULL"
    elif resolved_task_id:
        where += " AND task_id = ?"
        params.append(resolved_task_id)

    async with db.execute(  # type: ignore[attr-defined]
        "SELECT title, content FROM workspace_reports"
        f" WHERE {where}"
        " ORDER BY created_at",
        tuple(params),
    ) as cur:
        rows = await cur.fetchall()

    if not rows:
        raise FileNotFoundError(f"无已完成的报告：{ws_id}")

    docs = []
    for row in rows:
        title = row["title"] or "report"
        name = title if title.endswith(".md") else f"{title}.md"
        docs.append(_ReportDoc(name=name, content=row["content"] or ""))
    if legacy_only:
        suffix = "legacy"
    elif resolved_task_id:
        suffix = str(resolved_task_id)[:8]
    else:
        suffix = ws_id[:8]
    return _dispatch(docs, f"workspace-{ws_id[:8]}-{suffix}", fmt)


# ---------------------------------------------------------------------------
# Public: workspace chat export
# ---------------------------------------------------------------------------

async def export_workspace_chat(
    ws_id: str, ws_name: str, db: object
) -> tuple[bytes, str, str]:
    """Export workspace chat history as a Markdown file.

    Raises:
        FileNotFoundError: No chat messages for this workspace.
    """
    async with db.execute(  # type: ignore[attr-defined]
        "SELECT mode, role, content, created_at FROM workspace_chats"
        " WHERE workspace_id = ? ORDER BY created_at ASC",
        (ws_id,),
    ) as cur:
        rows = await cur.fetchall()

    if not rows:
        raise FileNotFoundError(f"无对话记录：{ws_id}")

    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    lines = [
        f"# 工作空间对话记录 — {ws_name}",
        "",
        f"导出时间：{now}",
        "",
        "---",
        "",
    ]

    MODE_LABELS = {"targeted": "结构化分析", "freeqa": "自由问答"}
    ROLE_LABELS = {"user": "用户", "assistant": "AI"}

    for row in rows:
        mode_label = MODE_LABELS.get(row["mode"], row["mode"])
        role_label = ROLE_LABELS.get(row["role"], row["role"])
        ts = row["created_at"][:16].replace("T", " ")
        lines.append(f"## [{mode_label}] {role_label} ({ts})")
        lines.append("")
        lines.append(redact_agent_diagnostic_text(row["content"] or ""))
        lines.append("")
        lines.append("---")
        lines.append("")

    md_text = "\n".join(lines)
    filename = f"chat-{ws_id[:8]}.md"
    return md_text.encode("utf-8"), filename, "text/markdown; charset=utf-8"


# ---------------------------------------------------------------------------
# Private helpers
# ---------------------------------------------------------------------------

def _dispatch(docs: list[_ReportDoc], prefix: str, fmt: str) -> tuple[bytes, str, str]:
    if fmt == "md":
        return _export_md_zip(docs, prefix)
    if fmt == "docx":
        return _export_docx(docs, prefix)
    if fmt == "xml":
        return _export_xml(docs, prefix)
    raise ValueError(f"不支持的导出格式: {fmt}")


def _export_md_zip(docs: list[_ReportDoc], prefix: str) -> tuple[bytes, str, str]:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for doc in docs:
            zf.writestr(doc.name, doc.content)
    return buf.getvalue(), f"{prefix}.zip", "application/zip"


def _export_docx(docs: list[_ReportDoc], prefix: str) -> tuple[bytes, str, str]:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:  # pragma: no cover
        raise RuntimeError("python-docx 未安装，请运行: pip install python-docx")

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(10.5)

    for report in docs:
        content = report.content
        if content.startswith("---"):
            end_idx = content.find("---", 3)
            if end_idx != -1:
                content = content[end_idx + 3:].strip()

        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("- "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped.startswith("| "):
                doc.add_paragraph(stripped, style="Normal")
            else:
                doc.add_paragraph(stripped)

        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return (
        buf.getvalue(),
        f"{prefix}.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _export_xml(docs: list[_ReportDoc], prefix: str) -> tuple[bytes, str, str]:
    import xml.etree.ElementTree as ET

    root = ET.Element("codetalk-reports")
    root.set("prefix", prefix)

    for doc in docs:
        content = doc.content
        report_el = ET.SubElement(root, "report")
        report_el.set("filename", doc.name)

        if content.startswith("---"):
            end_idx = content.find("---", 3)
            if end_idx != -1:
                frontmatter = content[3:end_idx].strip()
                content = content[end_idx + 3:].strip()
                meta_el = ET.SubElement(report_el, "metadata")
                for line in frontmatter.split("\n"):
                    if ": " in line:
                        key, val = line.split(": ", 1)
                        field_el = ET.SubElement(meta_el, key.strip())
                        field_el.text = val.strip()

        body_el = ET.SubElement(report_el, "content")
        body_el.text = content

    tree = ET.ElementTree(root)
    buf = io.StringIO()
    tree.write(buf, encoding="unicode", xml_declaration=True)
    return buf.getvalue().encode("utf-8"), f"{prefix}.xml", "application/xml"
