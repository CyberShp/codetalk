"""Deterministic parsers and provenance locators for F011 knowledge imports."""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
import re
from typing import Any


MAX_SOURCE_BYTES = 16 * 1024 * 1024
MAX_EXTRACTED_CHARACTERS = 2_000_000
_URL_RE = re.compile(r"https?://[^\s)\]>]+")

@dataclass(frozen=True)
class ParsedKnowledgeSource:
    status: str
    text: str
    locators: list[dict[str, Any]]
    parser: str
    error: str = ""


def parse_paste(text: str, *, keep_links: bool = False) -> ParsedKnowledgeSource:
    raw = _clean_text(str(text), keep_links=keep_links)
    if len(raw) > MAX_EXTRACTED_CHARACTERS:
        return ParsedKnowledgeSource("failed", "", [], "paste", "extracted_text_too_large")
    locators = [
        {"kind": "line", "start": number, "end": number, "excerpt": line}
        for number, line in enumerate(raw.splitlines(), start=1)
        if line
    ]
    return ParsedKnowledgeSource("parsed", raw, locators, "paste")


def parse_bytes(data: bytes, *, filename: str, keep_links: bool = False) -> ParsedKnowledgeSource:
    extension = Path(filename).suffix.lower()
    parser = {".docx": "python-docx", ".pdf": "pypdf", ".xlsx": "openpyxl"}.get(extension, "unsupported")
    if len(data) > MAX_SOURCE_BYTES:
        return ParsedKnowledgeSource("failed", "", [], parser, "source_too_large")
    try:
        if extension == ".docx":
            return _parse_docx(data, keep_links=keep_links)
        if extension == ".pdf":
            return _parse_pdf(data, keep_links=keep_links)
        if extension == ".xlsx":
            return _parse_xlsx(data, keep_links=keep_links)
        if extension in {".txt", ".md", ".log", ".csv"}:
            return parse_paste(data.decode("utf-8", errors="replace"), keep_links=keep_links)
        return ParsedKnowledgeSource("failed", "", [], parser, f"unsupported source format: {extension or 'unknown'}")
    except Exception as exc:
        return ParsedKnowledgeSource("failed", "", [], parser, f"{type(exc).__name__}: {exc}")


def _parse_docx(data: bytes, *, keep_links: bool) -> ParsedKnowledgeSource:
    from docx import Document

    document = Document(BytesIO(data))
    parts: list[str] = []
    locators: list[dict[str, Any]] = []
    headings: list[str] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = _clean_text(paragraph.text.strip(), keep_links=keep_links)
        style_name = str(paragraph.style.name or "")
        heading_match = re.fullmatch(r"Heading\s+(\d+)", style_name, flags=re.IGNORECASE)
        if heading_match and text:
            level = int(heading_match.group(1))
            headings = headings[: level - 1]
            headings.append(text)
        if text:
            parts.append(text)
            locators.append({"kind": "paragraph", "paragraph": index, "heading_path": list(headings), "excerpt": text})
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            for column_index, cell in enumerate(row.cells, start=1):
                text = _clean_text(cell.text.strip(), keep_links=keep_links)
                if text:
                    cell_name = f"{_column_name(column_index)}{row_index}"
                    parts.append(text)
                    locators.append({"kind": "table-cell", "table": table_index, "row": row_index, "column": column_index, "cell": cell_name, "excerpt": text})
    return _parsed_or_too_large("\n".join(parts), locators, "python-docx")


def _parse_pdf(data: bytes, *, keep_links: bool) -> ParsedKnowledgeSource:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(data))
    parts: list[str] = []
    locators: list[dict[str, Any]] = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = _clean_text((page.extract_text() or "").strip(), keep_links=keep_links)
        if text:
            parts.append(text)
            locators.append({"kind": "page-text-run", "page": page_number, "text_run": 1, "excerpt": text})
    if not parts:
        return ParsedKnowledgeSource("needs_ocr", "", [], "pypdf", "pdf_has_no_extractable_text")
    return _parsed_or_too_large("\n".join(parts), locators, "pypdf")


def _parse_xlsx(data: bytes, *, keep_links: bool) -> ParsedKnowledgeSource:
    from openpyxl import load_workbook

    workbook = load_workbook(BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    locators: list[dict[str, Any]] = []
    for worksheet in workbook.worksheets:
        headers: dict[int, str] = {}
        for row in worksheet.iter_rows():
            for cell in row:
                if cell.value is None:
                    continue
                text = _clean_text(str(cell.value), keep_links=keep_links)
                if cell.row == 1:
                    headers[cell.column] = text
                parts.append(text)
                locator = {"kind": "sheet-cell", "sheet": worksheet.title, "row": cell.row, "column": cell.column, "cell": cell.coordinate, "excerpt": text}
                if cell.row > 1 and headers.get(cell.column):
                    locator["header"] = headers[cell.column]
                    locator["header_cell"] = f"{_column_name(cell.column)}1"
                locators.append(locator)
    return _parsed_or_too_large("\n".join(parts), locators, "openpyxl")


def _column_name(index: int) -> str:
    value = ""
    while index:
        index, remainder = divmod(index - 1, 26)
        value = chr(65 + remainder) + value
    return value


def _clean_text(value: str, *, keep_links: bool) -> str:
    return value if keep_links else _URL_RE.sub("", value)


def _parsed_or_too_large(text: str, locators: list[dict[str, Any]], parser: str) -> ParsedKnowledgeSource:
    if len(text) > MAX_EXTRACTED_CHARACTERS:
        return ParsedKnowledgeSource("failed", "", [], parser, "extracted_text_too_large")
    return ParsedKnowledgeSource("parsed", text, locators, parser)
