"""Input ingestion for workbench task runs."""

from __future__ import annotations

import hashlib
import json
import shutil
from pathlib import Path
from typing import Any


TEXT_EXTS = frozenset({
    ".md", ".txt", ".patch", ".diff", ".json", ".xml", ".lcov", ".info",
    ".csv", ".tsv", ".yaml", ".yml", ".ini", ".cfg", ".conf", ".log",
})
DOCX_EXTS = frozenset({".docx"})
PDF_EXTS = frozenset({".pdf"})


def ingest_workbench_inputs(
    *,
    input_definitions: list[dict[str, Any]],
    inputs: dict[str, Any],
    artifact_dir: str | Path,
) -> dict[str, Any]:
    root = Path(artifact_dir) / "inputs"
    root.mkdir(parents=True, exist_ok=True)
    defs_by_id = {
        str(item.get("id")): item
        for item in input_definitions
        if isinstance(item, dict) and item.get("id")
    }
    snapshot: dict[str, Any] = {}
    for input_id, definition in defs_by_id.items():
        if bool(definition.get("required", False)) and _is_missing_input((inputs or {}).get(input_id)):
            raise ValueError(f"required input {input_id} is missing")
    for input_id, value in (inputs or {}).items():
        input_key = str(input_id)
        definition = defs_by_id.get(input_key, {})
        if definition and not bool(definition.get("required", False)) and _is_missing_input(value):
            continue
        input_type = str(definition.get("type") or "")
        if input_type in {"diff", "patch"} and _is_inline_patch_text(value):
            snapshot[input_key] = _ingest_inline_text_file(
                input_id=input_key,
                value=value,
                root=root,
                suffix=".patch",
            )
        elif input_type in {"file", "coverage_report", "diff", "patch"}:
            snapshot[input_key] = _ingest_file(
                input_id=input_key,
                value=value,
                root=root,
            )
        elif input_type == "file_set":
            snapshot[input_key] = _ingest_file_set(
                input_id=input_key,
                value=value,
                root=root,
            )
        else:
            snapshot[input_key] = value
        schema_errors = _validate_input_schema(
            value=snapshot[input_key],
            schema=definition.get("schema") or definition.get("json_schema"),
        )
        if schema_errors:
            raise ValueError(
                "input {} schema_validation_failed: {}".format(
                    input_key,
                    "; ".join(schema_errors),
                )
            )
    return snapshot


def _is_missing_input(value: Any) -> bool:
    if value is None:
        return True
    if isinstance(value, str):
        return not value.strip()
    if isinstance(value, (list, tuple, dict, set)):
        return len(value) == 0
    return False


def _ingest_file_set(*, input_id: str, value: Any, root: Path) -> dict[str, Any]:
    if not isinstance(value, (list, tuple)):
        raise ValueError(f"file_set input {input_id} must be a list of file paths")
    input_root = root / _safe_name(input_id)
    files: list[dict[str, Any]] = []
    for index, item in enumerate(value):
        file_info = _ingest_file(
            input_id=f"{input_id}_{index + 1}",
            value=item,
            root=input_root,
        )
        file_info["file_set_index"] = index
        files.append(file_info)
    if not files:
        raise ValueError(f"file_set input {input_id} is missing files")
    manifest = {
        "kind": "file_set",
        "input_id": input_id,
        "count": len(files),
        "files": files,
    }
    manifest_path = input_root / "file_set_manifest.json"
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    manifest["manifest_path"] = str(manifest_path)
    return manifest


def _ingest_file(*, input_id: str, value: Any, root: Path) -> dict[str, Any]:
    path_text = str(value.get("path") if isinstance(value, dict) else value or "").strip()
    if not path_text:
        raise ValueError(f"file input {input_id} is missing path")
    source = Path(path_text)
    if not source.exists() or not source.is_file():
        raise FileNotFoundError(path_text)
    input_root = root / _safe_name(input_id)
    original_dir = input_root / "original"
    original_dir.mkdir(parents=True, exist_ok=True)
    copied = original_dir / source.name
    shutil.copy2(source, copied)
    data = copied.read_bytes()
    sha256 = hashlib.sha256(data).hexdigest()

    parsed_text, parse_warnings = _extract_text(copied)
    parsed_text_path = input_root / "parsed_text.txt"
    parsed_text_path.write_text(parsed_text, encoding="utf-8")
    chunks = _chunks(parsed_text)
    chunks_path = input_root / "chunks.json"
    chunks_path.write_text(json.dumps(chunks, ensure_ascii=False, indent=2), encoding="utf-8")
    if not parse_warnings and not parsed_text:
        parse_warnings = ["text_extraction_empty"]
    metadata = {
        "kind": "file",
        "input_id": input_id,
        "original_path": str(source),
        "copied_path": str(copied),
        "filename": source.name,
        "suffix": source.suffix.lower(),
        "size_bytes": len(data),
        "sha256": sha256,
        "parsed_text_path": str(parsed_text_path),
        "chunks_path": str(chunks_path),
        "parse_warnings": parse_warnings,
    }
    metadata_path = input_root / "file_metadata.json"
    metadata_path.write_text(json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8")
    metadata["metadata_path"] = str(metadata_path)
    return metadata


def _ingest_inline_text_file(
    *,
    input_id: str,
    value: Any,
    root: Path,
    suffix: str,
) -> dict[str, Any]:
    text = _inline_text_value(value)
    if not text:
        raise ValueError(f"file input {input_id} is missing inline text")
    input_root = root / _safe_name(input_id)
    original_dir = input_root / "original"
    original_dir.mkdir(parents=True, exist_ok=True)
    filename = _safe_name(input_id) + suffix
    copied = original_dir / filename
    copied.write_text(text, encoding="utf-8")
    data = copied.read_bytes()
    sha256 = hashlib.sha256(data).hexdigest()
    parsed_text, parse_warnings = _extract_text(copied)
    parsed_text_path = input_root / "parsed_text.txt"
    parsed_text_path.write_text(parsed_text, encoding="utf-8")
    chunks = _chunks(parsed_text)
    chunks_path = input_root / "chunks.json"
    chunks_path.write_text(json.dumps(chunks, ensure_ascii=False, indent=2), encoding="utf-8")
    metadata = {
        "kind": "file",
        "input_id": input_id,
        "original_path": "",
        "copied_path": str(copied),
        "filename": filename,
        "suffix": suffix,
        "size_bytes": len(data),
        "sha256": sha256,
        "parsed_text_path": str(parsed_text_path),
        "chunks_path": str(chunks_path),
        "parse_warnings": parse_warnings,
        "inline_text": True,
    }
    metadata_path = input_root / "file_metadata.json"
    metadata_path.write_text(json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8")
    metadata["metadata_path"] = str(metadata_path)
    return metadata


def _is_inline_patch_text(value: Any) -> bool:
    text = _inline_text_value(value)
    return "diff --git " in text or ("\n--- " in text and "\n+++ " in text)


def _inline_text_value(value: Any) -> str:
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, dict):
        return str(value.get("text") or value.get("content") or "").strip()
    return ""


def _extract_text(path: Path) -> tuple[str, list[str]]:
    suffix = path.suffix.lower()
    if suffix in TEXT_EXTS:
        return path.read_text(encoding="utf-8", errors="replace"), []
    if suffix in DOCX_EXTS:
        return _extract_docx_text(path)
    if suffix in PDF_EXTS:
        return _extract_pdf_text(path)
    return "", ["unsupported_file_type_for_text_extraction"]


def _extract_docx_text(path: Path) -> tuple[str, list[str]]:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - dependency is required in backend requirements
        return "", [f"docx_extraction_unavailable:{type(exc).__name__}"]
    try:
        document = Document(str(path))
    except Exception as exc:
        return "", [f"docx_extraction_failed:{type(exc).__name__}"]
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    table_cells: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    table_cells.append(cell.text)
    text = "\n".join([*paragraphs, *table_cells]).strip()
    return text, [] if text else ["text_extraction_empty"]


def _extract_pdf_text(path: Path) -> tuple[str, list[str]]:
    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]
    except Exception:
        return "", ["pdf_extraction_unavailable:pypdf_not_installed"]
    try:
        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        return "", [f"pdf_extraction_failed:{type(exc).__name__}"]
    text = "\n".join(page for page in pages if page).strip()
    return text, [] if text else ["text_extraction_empty"]


def _chunks(text: str, *, chunk_size: int = 4000) -> list[dict[str, Any]]:
    if not text:
        return []
    chunks: list[dict[str, Any]] = []
    for index, start in enumerate(range(0, len(text), chunk_size)):
        chunk = text[start:start + chunk_size]
        chunks.append({
            "chunk_index": index,
            "start_char": start,
            "end_char": start + len(chunk),
            "content": chunk,
        })
    return chunks


def _safe_name(value: str) -> str:
    return "".join(char if char.isalnum() or char in {"-", "_"} else "_" for char in value) or "input"


def _validate_input_schema(*, value: Any, schema: Any) -> list[str]:
    if not isinstance(schema, dict):
        return []
    return _validate_schema_fragment(value, schema)


def _validate_schema_fragment(value: Any, schema: dict[str, Any], *, path: str = "$") -> list[str]:
    errors: list[str] = []
    expected_type = str(schema.get("type") or "").strip()
    if expected_type and not _matches_schema_type(value, expected_type):
        errors.append(f"{path} expected {expected_type}")
        return errors

    enum = schema.get("enum")
    if isinstance(enum, list) and value not in enum:
        errors.append(f"{path} must be one of: {', '.join(str(item) for item in enum)}")

    if isinstance(value, str):
        min_length = schema.get("minLength")
        if isinstance(min_length, int) and len(value) < min_length:
            errors.append(f"{path} length must be >= {min_length}")

    if isinstance(value, dict):
        for field in schema.get("required") or []:
            if isinstance(field, str) and field not in value:
                errors.append(f"missing required field: {field}")
        properties = schema.get("properties") or {}
        if isinstance(properties, dict):
            for field_name, property_schema in properties.items():
                if field_name not in value or not isinstance(property_schema, dict):
                    continue
                errors.extend(
                    _validate_schema_fragment(
                        value[field_name],
                        property_schema,
                        path=f"{path}.{field_name}",
                    )
                )
    return errors


def _matches_schema_type(value: Any, expected_type: str) -> bool:
    if expected_type == "object":
        return isinstance(value, dict)
    if expected_type == "array":
        return isinstance(value, list)
    if expected_type == "string":
        return isinstance(value, str)
    if expected_type == "number":
        return isinstance(value, (int, float)) and not isinstance(value, bool)
    if expected_type == "integer":
        return isinstance(value, int) and not isinstance(value, bool)
    if expected_type == "boolean":
        return isinstance(value, bool)
    if expected_type == "null":
        return value is None
    return True
